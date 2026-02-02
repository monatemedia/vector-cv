"""
Generate Word documents from CV markdown
Converts markdown-formatted CVs to professional Word documents with proper formatting
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import json


def json_cv_to_markdown(cv_data):
    """Convert structured CV JSON to markdown format"""
    if isinstance(cv_data, str):
        try:
            cv_data = json.loads(cv_data)
        except json.JSONDecodeError:
            return cv_data
    
    if not isinstance(cv_data, dict) or 'header' not in cv_data:
        return str(cv_data)
    
    md_lines = []
    
    # Header
    header = cv_data.get('header', {})
    md_lines.append(f"# {header.get('name', 'Edward Baitsewe')}")
    md_lines.append(f"**{header.get('title', 'Full Stack Developer')}**")
    md_lines.append(f"📍 {header.get('location', '')} | 📞 {header.get('phone', '')} | 📧 {header.get('email', '')}")
    
    contact_links = []
    if header.get('linkedin'):
        contact_links.append(f"🔗 {header['linkedin']}")
    if header.get('portfolio'):
        contact_links.append(f"🌐 {header['portfolio']}")
    if header.get('github'):
        contact_links.append(f"🐙 {header['github']}")
    
    if contact_links:
        md_lines.append(" | ".join(contact_links))
    
    md_lines.append("")
    md_lines.append("## 🔹 Summary")
    md_lines.append(cv_data.get('summary', ''))
    md_lines.append("")
    md_lines.append("---")
    md_lines.append("")
    md_lines.append("## 🔹 Core Technical Strengths")
    
    tech_strengths = cv_data.get('technical_strengths', {})
    for category, skills in tech_strengths.items():
        md_lines.append(f"* **{category}:** {skills}")
    
    md_lines.append("")
    md_lines.append("---")
    md_lines.append("")
    md_lines.append("## 🔹 Key Projects")
    md_lines.append("")
    
    key_projects = cv_data.get('key_projects', [])
    source_chunks = cv_data.get('_source_chunks', {})
    
    for project in key_projects:
        title = project.get('title', '')
        content = project.get('content', '')
        demo_table = project.get('demo_table')
        
        md_lines.append(f"**{title}**")
        md_lines.append("")
        md_lines.append(content)
        md_lines.append("")
        
        if demo_table:
            md_lines.append("**Demo Available:**")
            md_lines.append("")
            md_lines.append("| Field | Value |")
            md_lines.append("|-------|-------|")
            for field, value in demo_table.items():
                md_lines.append(f"| {field} | `{value}` |")
            md_lines.append("")
        
        source_chunk = source_chunks.get(title)
        if source_chunk:
            links = extract_links_from_source(source_chunk.get('content', ''))
            if links:
                formatted = format_links_as_pipe_separated(links)
                md_lines.append(formatted)
                md_lines.append("")
    
    md_lines.append("---")
    md_lines.append("")
    md_lines.append("## 🔹 Professional Experience")
    md_lines.append("")
    md_lines.append(cv_data.get('professional_experience', ''))
    md_lines.append("")
    md_lines.append("---")
    md_lines.append("")
    md_lines.append("## 🔹 Education")
    md_lines.append("")
    md_lines.append(cv_data.get('education', ''))
    
    return "\n".join(md_lines)


def extract_links_from_source(content):
    """Extract markdown links from source content"""
    link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    return re.findall(link_pattern, content)


def format_links_as_pipe_separated(links):
    """Format links as 'Label: URL | Label: URL'"""
    formatted = [f"{label}: {url}" for label, url in links]
    return " | ".join(formatted)


def parse_inline_markdown(text):
    """
    Parse inline markdown formatting (bold, italic, code) and return segments
    Returns list of dicts with 'text', 'bold', 'italic', 'code' properties
    """
    segments = []
    i = 0
    
    while i < len(text):
        # Check for **bold**
        if i < len(text) - 1 and text[i:i+2] == '**':
            end = text.find('**', i + 2)
            if end != -1:
                segments.append({'text': text[i+2:end], 'bold': True, 'italic': False, 'code': False})
                i = end + 2
                continue
        
        # Check for *italic*
        if text[i] == '*':
            end = text.find('*', i + 1)
            if end != -1:
                segments.append({'text': text[i+1:end], 'bold': False, 'italic': True, 'code': False})
                i = end + 1
                continue
        
        # Check for `code`
        if text[i] == '`':
            end = text.find('`', i + 1)
            if end != -1:
                segments.append({'text': text[i+1:end], 'bold': False, 'italic': False, 'code': True})
                i = end + 1
                continue
        
        # Regular text - find next special character
        next_special = len(text)
        for special in ['**', '*', '`']:
            pos = text.find(special, i)
            if pos != -1 and pos < next_special:
                next_special = pos
        
        if next_special > i:
            segments.append({'text': text[i:next_special], 'bold': False, 'italic': False, 'code': False})
            i = next_special
        else:
            i += 1
    
    return segments if segments else [{'text': text, 'bold': False, 'italic': False, 'code': False}]


def add_formatted_run(paragraph, segment):
    """Add a formatted run to a paragraph"""
    run = paragraph.add_run(segment['text'])
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    
    if segment.get('bold'):
        run.font.bold = True
    if segment.get('italic'):
        run.font.italic = True
    if segment.get('code'):
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(68, 68, 68)


def create_table(doc, rows):
    """Create a properly formatted table"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Light Grid Accent 1'
    
    for i, (field, value) in enumerate(rows):
        # Field cell
        cell1 = table.rows[i].cells[0]
        cell1.text = field
        cell1.paragraphs[0].runs[0].font.bold = True
        cell1.paragraphs[0].runs[0].font.size = Pt(10)
        
        # Value cell
        cell2 = table.rows[i].cells[1]
        cell2.text = value.strip('`')
        cell2.paragraphs[0].runs[0].font.name = 'Courier New'
        cell2.paragraphs[0].runs[0].font.size = Pt(10)
        cell2.paragraphs[0].runs[0].font.color.rgb = RGBColor(68, 68, 68)
    
    return table


def parse_markdown_to_docx(markdown_text: str, output_path: str):
    """Convert markdown CV to Word document with proper formatting"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    lines = markdown_text.split('\n')
    i = 0
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # H1 - Name
        if line.startswith('# '):
            name = line[2:].strip()
            p = doc.add_paragraph()
            run = p.add_run(name)
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(31, 78, 121)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Job title (bold line after name)
        if line.startswith('**') and line.endswith('**') and '|' not in line and ':' not in line:
            title = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(68, 68, 68)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Contact info (lines with emoji)
        if any(emoji in line for emoji in ['📍', '📞', '📧', '🔗', '🌐', '🐙']):
            clean_line = re.sub(r'[📍📞📧🔗🌐🐙]', '', line)
            clean_line = clean_line.replace('|', ' • ')
            
            p = doc.add_paragraph()
            run = p.add_run(clean_line.strip())
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(102, 102, 102)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # H2 - Section headers
        if line.startswith('## '):
            section = line[2:].strip().replace('🔹', '').strip()
            p = doc.add_heading(section, level=2)
            p.runs[0].font.color.rgb = RGBColor(31, 78, 121)
            p.runs[0].font.size = Pt(14)
            i += 1
            continue
        
        # Horizontal rule - just skip, don't add space
        if line == '---':
            i += 1
            continue
        
        # Table detection
        if line.startswith('|') and not in_table:
            # Check if next line is separator
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('|---'):
                in_table = True
                table_rows = []
                i += 2  # Skip header and separator
                continue
        
        # Collect table rows
        if in_table:
            if line.startswith('|') and '---' not in line:
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if len(cells) == 2:
                    table_rows.append((cells[0], cells[1]))
                i += 1
                continue
            else:
                # End of table
                if table_rows:
                    create_table(doc, table_rows)
                in_table = False
                # Don't increment i, process this line normally
                continue
        
        # Bullet points
        if line.startswith('* '):
            content = line[2:].strip()
            segments = parse_inline_markdown(content)
            
            p = doc.add_paragraph(style='List Bullet')
            for segment in segments:
                add_formatted_run(p, segment)
            i += 1
            continue
        
        # Project links (pipe-separated with http)
        if '|' in line and ': http' in line and not line.startswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(11)  # Same size as body text
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
            i += 1
            continue
        
        # Regular paragraphs with inline formatting
        segments = parse_inline_markdown(line)
        p = doc.add_paragraph()
        
        for segment in segments:
            add_formatted_run(p, segment)
        
        i += 1
    
    doc.save(output_path)
    return output_path


def generate_cv_docx(cv_data: str, company_name: str, job_title: str, output_dir: str = "./") -> str:
    """Generate Word document from CV data (JSON or markdown)"""
    markdown = json_cv_to_markdown(cv_data)
    filename = f"CV_{company_name}_{job_title}.docx".replace(' ', '_').replace('/', '_')
    output_path = f"{output_dir}/{filename}"
    parse_markdown_to_docx(markdown, output_path)
    return output_path


def generate_cv_markdown(cv_data: str, company_name: str, job_title: str, output_dir: str = "./") -> str:
    """Generate Markdown file from CV data (JSON or markdown)"""
    markdown = json_cv_to_markdown(cv_data)
    filename = f"CV_{company_name}_{job_title}.md".replace(' ', '_').replace('/', '_')
    output_path = f"{output_dir}/{filename}"
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown)
    
    return output_path


def generate_cover_letter_docx(cover_letter_markdown: str, company_name: str, job_title: str, output_dir: str = "./") -> str:
    """Generate Word document from cover letter markdown"""
    filename = f"CoverLetter_{company_name}_{job_title}.docx".replace(' ', '_').replace('/', '_')
    output_path = f"{output_dir}/{filename}"

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = cover_letter_markdown.split('\n')

    for line in lines:
        line = line.strip()

        if not line or line.startswith('# ') or line.startswith('🔹'):
            continue

        # To/Subject lines
        if line.startswith('**To:**') or line.startswith('**Subject:**'):
            segments = parse_inline_markdown(line)
            p = doc.add_paragraph()
            for segment in segments:
                add_formatted_run(p, segment)
            continue

        # Greeting
        if line.startswith('Dear '):
            doc.add_paragraph()
            p = doc.add_paragraph(line)
            continue

        # Closing
        if line in ['Best regards,', 'Sincerely,', 'Kind regards,']:
            doc.add_paragraph()
            p = doc.add_paragraph(line)
            continue

        # Contact info
        if any(char in line for char in ['+27', '@']):
            doc.add_paragraph(line)
            continue

        # Regular paragraphs
        segments = parse_inline_markdown(line)
        p = doc.add_paragraph()
        for segment in segments:
            add_formatted_run(p, segment)

    doc.save(output_path)
    return output_path


def generate_cover_letter_markdown(cover_letter_markdown: str, company_name: str, job_title: str, output_dir: str = "./") -> str:
    """Generate Markdown file from cover letter markdown"""
    filename = f"CoverLetter_{company_name}_{job_title}.md".replace(' ', '_').replace('/', '_')
    output_path = f"{output_dir}/{filename}"
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(cover_letter_markdown)
    
    return output_path


if __name__ == "__main__":
    sample_json = {
        "header": {
            "name": "Edward Baitsewe",
            "title": "Full Stack Developer",
            "location": "Parow, Cape Town",
            "phone": "+27 78 324 5326",
            "email": "edward@monatemedia.com",
            "linkedin": "https://www.linkedin.com/in/edwardbaitsewe",
            "portfolio": "https://monatemedia.com/portfolio",
            "github": "https://github.com/monatemedia"
        },
        "summary": "Full stack developer with 5 years of experience building scalable **Laravel** applications.",
        "technical_strengths": {
            "Backend": "PHP (Laravel 9-12), Python",
            "Frontend": "JavaScript, React"
        },
        "key_projects": [
            {
                "title": "ActuallyFind – Core Platform",
                "content": "**Production Marketplace:** Built vehicle marketplace using **Laravel 12**.\n**Tech Stack:** **Docker**, **PostgreSQL**",
                "demo_table": {
                    "URL": "https://dealership.monatemedia.com/",
                    "Email": "user@example.com",
                    "Password": "password"
                }
            }
        ],
        "professional_experience": "**Full Stack Developer** | *Monate Media* | 2021 - Present\n\n* Engineered **Akaunting** ERP solutions",
        "education": "* **Postgraduate Diploma** – University of the Free State",
        "_source_chunks": {
            "ActuallyFind – Core Platform": {
                "content": "[Demo](https://dealership.monatemedia.com/) [Production](https://actuallyfind.com/) [GitHub](https://github.com/monatemedia/dealership)"
            }
        }
    }
    
    print("Testing JSON to Markdown conversion...")
    markdown = json_cv_to_markdown(sample_json)
    print(markdown)
    print("\n" + "="*80 + "\n")
    
    print("Generating test files...")
    generate_cv_docx(json.dumps(sample_json), "TestCompany", "Developer", "./")
    generate_cv_markdown(json.dumps(sample_json), "TestCompany", "Developer", "./")
    print("✅ Test files generated:")
    print("   - CV_TestCompany_Developer.docx")
    print("   - CV_TestCompany_Developer.md")